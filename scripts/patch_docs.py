#!/usr/bin/env python3
"""根据项目实际范围修改docx文档内容，不改变格式"""

from docx import Document
import os

DOCS = "/home/xjh01/ShangHaiFoodData/docs"


def replace_in_para(para, old, new):
    """把段落所有run的文本合并后替换，结果放回第一个run，保留段落格式"""
    full = "".join(r.text for r in para.runs)
    if old in full:
        replaced = full.replace(old, new)
        if para.runs:
            para.runs[0].text = replaced
            for r in para.runs[1:]:
                r.text = ""
        return True
    return False


def patch(doc, pairs):
    n = 0
    for para in doc.paragraphs:
        for old, new in pairs:
            if replace_in_para(para, old, new):
                n += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in pairs:
                        if replace_in_para(para, old, new):
                            n += 1
    return n


# ─────────────────────────────────────────────
# 文档2：软件需求规约
# ─────────────────────────────────────────────
p2 = os.path.join(DOCS, "2.第6组_软件需求规约.docx")
d2 = Document(p2)
r2 = [
    # 模块总数描述（基于当前文档实际内容）
    (
        "系统由四大核心模块组成，分别为：用户登录与认证模块、数据采集模块、数据分析模块、数据可视化模块。",
        "系统由八大核心模块组成，分别为：用户登录与认证模块、数据采集模块、数据导入模块（支持MySQL/HDFS/Hive）、数据分析模块、数据可视化模块、机器学习评分预测模块、可视化报告导出模块、数据清洗日志模块。",
    ),
    # 模块列表简述
    (
        "系统分成四大模块：用户登录模块、数据采集模块、数据分析模块、数据可视化模块",
        "系统分成八大模块：用户登录模块、数据采集模块、数据导入模块、数据分析模块、数据可视化模块、机器学习预测模块、报告导出模块、数据清洗日志模块",
    ),
    # 核心需求第五条（当前文档已经是简化后的版本）
    (
        "第五：所有数据修改、管理操作均需在登录状态下进行，基于角色实现权限分级管控。",
        "第五：要实现数据导入功能，支持将数据集（CSV/JSON/Excel）导入到MySQL数据库、上传至HDFS分布式存储或加载到Hive数据仓库表，每个清洗步骤均记录操作日志，供前端实时查看。"
        "第六：要实现可视化报告导出功能，后端使用matplotlib生成图表图片、reportlab排版组合生成PDF，包含各图表图片及对应文字说明，支持浏览器一键下载。"
        "第七：要实现机器学习评分预测功能，基于随机森林回归算法（scikit-learn），根据菜系、区域、人均消费预测餐厅评分，并展示R²、RMSE精度指标及置信区间。"
        "第八：要实现数据清洗过程日志显示功能，前端页面实时展示每步清洗操作的阶段（load/dedup/normalize/validate/insert/hdfs/hive）、级别（INFO/WARNING/ERROR）和影响记录数，支持按任务和级别筛选，支持自动刷新。"
        "第九：所有数据修改、管理操作均需在登录状态下进行，基于角色实现权限分级管控。",
    ),
    # 数据库描述（新增大数据组件）
    (
        "数据库：MySQL 8.0",
        "数据库：MySQL 8.0（主要结构化存储），可选扩展：Hadoop HDFS（分布式文件存储）、Hive（数据仓库）",
    ),
    # 用户角色描述
    (
        "普通用户（user），可以进行如下操作：查看餐厅信息、查看分析报告、收藏餐厅、获取个性化推荐等。",
        "普通用户（user），可以进行如下操作：搜索餐厅、查看分析图表、使用机器学习评分预测、下载PDF分析报告、查看数据清洗日志等。",
    ),
]
n2 = patch(d2, r2)
d2.save(p2)
print(f"文档2 完成 {n2} 处替换")


# ─────────────────────────────────────────────
# 文档4：数据存储设计说明书
# ─────────────────────────────────────────────
p4 = os.path.join(DOCS, "4.第6组_数据存储设计说明书.docx")
d4 = Document(p4)
r4 = [
    # 存储技术选型——新增 HDFS/Hive
    (
        "结构化数据库：MySQL 8.0，适用于存储餐厅信息、评论、用户等关系型数据，支持事务处理、复杂查询及索引优化。",
        "结构化数据库：MySQL 8.0，适用于存储餐厅信息、评论、用户、数据清洗日志（clean_log）等关系型数据，支持事务处理、复杂查询及索引优化。"
        "分布式文件存储（可选）：Hadoop HDFS，通过WebHDFS REST API（端口9870）上传CSV格式数据集，供Hive加载使用。"
        "数据仓库（可选）：Hive，通过HiveServer2（端口10000）执行LOAD DATA INPATH将HDFS数据加载到restaurant_info表，支持大规模SQL分析。",
    ),
    # 存储架构总体描述
    (
        "本系统采用MySQL 8.0作为唯一结构化数据存储。",
        "本系统以MySQL 8.0为主要结构化数据存储，可选扩展至HDFS+Hive大数据存储体系。数据导入支持三种目标：直接写入MySQL、上传至HDFS、通过HiveServer2加载至Hive表。",
    ),
    # 主表数量更新（新增 clean_log）
    (
        "主表：4个（restaurant_info、restaurant_review、sys_user、crawl_task）",
        "主表：5个（restaurant_info、restaurant_review、sys_user、crawl_task、clean_log），其中 clean_log 表新增用于记录每次数据导入的清洗步骤日志。",
    ),
    # 非结构化数据库描述（改为 HDFS/Hive 说明）
    (
        "本系统不使用非结构化数据库，所有数据存储于MySQL 8.0。",
        "分布式存储（可选）：系统支持将清洗后的数据集以CSV格式上传至HDFS（通过WebHDFS REST API），并通过Hive执行LOAD DATA INPATH加载为可查询的外部表，满足大数据分析场景需求。",
    ),
    # clean_log 表在表列表中
    (
        "爬虫任务表（crawl_task）",
        "爬虫任务表（crawl_task）",  # 保持不变（在下面的报告任务表之前）
    ),
    (
        "报告任务表（report_task）",
        "数据清洗日志表（clean_log）——新增，字段：id、task_name（任务名）、stage（清洗阶段：load/dedup/normalize/validate/insert/hdfs/hive）、level（INFO/WARNING/ERROR）、message（日志内容）、record_count（影响记录数）、create_time（创建时间）",
    ),
    # 物理模型描述更新
    (
        "分析结果按需计算后直接返回，报告以 PDF 格式由后端实时生成并返回给客户端下载",
        "分析结果按需计算后直接返回；PDF报告由后端使用matplotlib+reportlab实时生成，含5张分析图表及对应文字说明；数据清洗日志持久化至clean_log表，供前端实时查询。",
    ),
    # RBAC 描述
    (
        "．普通用户（user），可以进行如下操作：查看餐厅信息、查看分析报告、收藏餐厅、获取个性化推荐等。",
        "．普通用户（user），可以进行如下操作：搜索餐厅、查看分析图表、使用ML评分预测、下载PDF报告、查看清洗日志等。",
    ),
    (
        "．平台管理员（admin），可以进行如下操作：用户管理、爬虫任务管理、查看所有数据等全部操作。",
        "．平台管理员（admin），可以进行如下操作：用户管理、数据导入（MySQL/HDFS/Hive）、查看数据清洗日志、下载分析报告、管理爬虫任务等全部操作。",
    ),
]
n4 = patch(d4, r4)
d4.save(p4)
print(f"文档4 完成 {n4} 处替换")


# ─────────────────────────────────────────────
# 文档5：系统架构与详细设计说明书
# ─────────────────────────────────────────────
p5 = os.path.join(DOCS, "5.第6组_系统架构与详细设计说明书.docx")
d5 = Document(p5)
r5 = [
    # 离线分析技术选型——加入 scikit-learn、reportlab、matplotlib
    (
        "技术选型：Pandas（数据处理）+ SQLAlchemy（ORM批量写入）。",
        "技术选型：Pandas（数据处理）+ SQLAlchemy（ORM批量写入）+ scikit-learn（随机森林评分预测）+ matplotlib（图表生成）+ reportlab（PDF报告排版）。",
    ),
    # 流处理——加入 clean_log
    (
        "技术选型：Scrapy Pipeline（实时清洗入库）。",
        "技术选型：Scrapy Pipeline（实时清洗入库）+ clean_log 表（记录清洗步骤日志：阶段/级别/影响记录数，供前端实时查询）。",
    ),
    # 数据处理流程——加入导入流程
    (
        "分析流程：Pandas加载MySQL数据 → 执行分析算法 → 结果直接返回给API。",
        "分析流程：Pandas加载MySQL数据 → 执行分析算法 → 结果直接返回给API。"
        "数据导入流程：上传CSV/JSON/Excel文件 → Pandas解析 → 数据清洗（去重/脱敏/标准化，每步写入clean_log）→ 写入MySQL/HDFS/Hive。"
        "报告生成流程：查询分析数据 → matplotlib生成5张图表（PNG）→ reportlab组合图表+文字说明 → 生成PDF → 浏览器下载。"
        "ML预测流程：从MySQL加载数据 → LabelEncoder编码分类特征 → 训练RandomForestRegressor（100棵树）→ 持久化模型 → 接受预测请求 → 返回评分+置信区间。",
    ),
    # 存储位置——加入HDFS/Hive
    (
        "存储位置：结构化数据存入MySQL 8.0（shanghai_food库）。",
        "存储位置：结构化数据存入MySQL 8.0（shanghai_food库）；可选扩展至HDFS（WebHDFS REST API上传，端口9870）和Hive（HiveServer2，端口10000，LOAD DATA INPATH）。ML模型持久化到本地文件（/tmp/shanghai_food_rf_model.pkl）。",
    ),
    # 缓存机制——加入ML模型缓存
    (
        "缓存机制：高频分析接口使用Python内存缓存（functools.lru_cache），减少重复计算。",
        "缓存机制：高频分析接口使用Python内存缓存（functools.lru_cache），减少重复计算；机器学习模型训练后序列化为pickle文件，后续预测直接加载，避免重复训练。",
    ),
    # 性能测试指标更新
    (
        "区域分析接口响应时间≤3s，分析结果同步返回。",
        "区域分析接口响应时间≤3s；PDF报告生成接口≤30s（matplotlib渲染+reportlab排版同步完成）；ML预测接口（模型已加载时）≤1s；数据导入接口支持批量提交，每100条提交一次事务。",
    ),
    # 模块划分——更新为8大模块
    (
        "系统划分为八大功能模块：数据采集模块、数据导入模块、餐厅信息管理模块、数据分析模块、机器学习预测模块、报告导出模块、数据清洗日志模块、后台管理模块",
        "系统划分为八大功能模块：①数据采集模块（美团Scrapy爬虫）②数据导入模块（CSV/JSON/Excel → MySQL/HDFS/Hive）③餐厅信息管理模块（搜索/详情）④数据分析模块（区域/菜系/价格/评分/均价）⑤机器学习预测模块（RandomForest评分预测）⑥可视化报告导出模块（matplotlib+reportlab PDF）⑦数据清洗日志模块（clean_log实时展示）⑧后台管理模块（用户管理）",
    ),
    # 数据安全——更新角色描述
    (
        "最小权限原则：RBAC四级角色（admin/user），每个角色仅拥有完成其业务所需的最小API权限，实现于backend/app/core/deps.py。",
        "最小权限原则：RBAC两级角色（admin/user），每个角色仅拥有完成其业务所需的最小API权限，实现于backend/app/core/deps.py。公开接口（分析/搜索/预测/报告）无需认证，管理接口需admin角色。",
    ),
    # 模型评估——更新推荐为ML
    (
        "模型评估：情感分析准确率目标≥80%，推荐系统召回率目标≥60%，通过测试数据集定期评估。",
        "模型评估：情感分析（SnowNLP）准确率目标≥80%；机器学习评分预测（RandomForestRegressor）R²目标≥0.6，RMSE目标≤0.5，通过train_test_split（8:2）验证。",
    ),
]
n5 = patch(d5, r5)
d5.save(p5)
print(f"文档5 完成 {n5} 处替换")
